#!/usr/bin/env python3
"""Convert ADAudit Plus feature page markdown files to .docx format."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import os


# ─── Helpers ────────────────────────────────────────────────────────────────

def add_hyperlink(para, text, url):
    """Add a hyperlink run to a paragraph."""
    part = para.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    r_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2E74B5')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    r_el.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r_el.append(t)
    hyperlink.append(r_el)
    para._p.append(hyperlink)


def process_inline(para, text):
    """Process inline markdown: **bold**, [anchor](url), plain text."""
    pattern = r'(\*\*[^*]+\*\*|\[[^\]]+\]\([^\)]+\))'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        bold_m = re.match(r'\*\*([^*]+)\*\*', part)
        link_m = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', part)
        if bold_m:
            run = para.add_run(bold_m.group(1))
            run.bold = True
        elif link_m:
            add_hyperlink(para, link_m.group(1), link_m.group(2))
        else:
            para.add_run(part)


def heading_color(doc, text, level):
    """Add a heading paragraph with ME brand colours."""
    p = doc.add_heading('', level=level)
    p.clear()
    process_inline(p, text)
    if p.runs:
        run = p.runs[0]
        run.bold = True
        if level == 1:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def parse_table_block(lines, start):
    """Return (list_of_row_lists, next_line_index)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # skip separator row
        if re.match(r'^\|[\-| :]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def set_cell_bg(cell, hex_color):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


# ─── Publish converter ───────────────────────────────────────────────────────

def build_publish_docx(md_path, out_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    with open(md_path, 'r', encoding='utf-8') as f:
        raw = f.read()

    # Strip draft header, keep META BLOCK → end of page content
    meta_start = raw.find('## META BLOCK')
    end_marker = '*End of draft'
    end_pos = raw.find(end_marker)
    if meta_start != -1:
        raw = raw[meta_start:end_pos if end_pos != -1 else len(raw)].rstrip()

    lines = raw.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule → skip
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # H1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading_color(doc, stripped[2:], 1)
            i += 1
            continue

        # H2
        if stripped.startswith('## ') and not stripped.startswith('### '):
            heading_color(doc, stripped[3:], 2)
            i += 1
            continue

        # H3
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            heading_color(doc, stripped[4:], 3)
            i += 1
            continue

        # Table
        if stripped.startswith('|'):
            rows, next_i = parse_table_block(lines, i)
            if rows:
                col_count = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=col_count)
                tbl.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx in range(col_count):
                        cell_text = row[c_idx] if c_idx < len(row) else ''
                        cell = tbl.rows[r_idx].cells[c_idx]
                        cell.text = ''
                        para = cell.paragraphs[0]
                        process_inline(para, cell_text)
                        if r_idx == 0:
                            for run in para.runs:
                                run.bold = True
                            set_cell_bg(cell, 'DAEEF3')
                doc.add_paragraph()
            i = next_i
            continue

        # Bullet
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, stripped[2:])
            i += 1
            continue

        # Screenshot placeholder
        if stripped.startswith('ADD SCREENSHOT HERE'):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.italic = True
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            i += 1
            continue

        # CTA line: standalone [text] or [text](url) with NOTHING after the closing bracket/paren.
        # If there is trailing text after the link, treat as a normal paragraph (not a CTA).
        _is_cta = False
        if re.match(r'^\[.+\]', stripped) and not stripped.startswith('!['):
            _link_only_m = re.match(r'^\[([^\]]+)\]\(([^\)]+)\)\s*$', stripped)
            _bracket_only_m = re.match(r'^\[([^\]]+)\]\s*$', stripped)
            if _link_only_m or _bracket_only_m:
                _is_cta = True
        if _is_cta:
            link_m = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', stripped)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if link_m:
                run = p.add_run(f'[ {link_m.group(1)} ]')
            else:
                txt = re.match(r'\[([^\]]+)\]', stripped).group(1)
                run = p.add_run(f'[ {txt} ]')
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            i += 1
            continue

        # Normal paragraph (may contain bold/links)
        p = doc.add_paragraph()
        process_inline(p, stripped)
        i += 1

    doc.save(out_path)
    print(f'  Saved: {out_path}')


# ─── Review converter ────────────────────────────────────────────────────────

def build_review_docx(md_path, out_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading_color(doc, stripped[2:], 1)
            i += 1
            continue

        if stripped.startswith('## ') and not stripped.startswith('### '):
            heading_color(doc, stripped[3:], 2)
            i += 1
            continue

        if stripped.startswith('### ') and not stripped.startswith('#### '):
            heading_color(doc, stripped[4:], 3)
            i += 1
            continue

        if stripped.startswith('#### '):
            heading_color(doc, stripped[5:], 3)
            i += 1
            continue

        if stripped.startswith('|'):
            rows, next_i = parse_table_block(lines, i)
            if rows:
                col_count = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=col_count)
                tbl.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx in range(col_count):
                        cell_text = row[c_idx] if c_idx < len(row) else ''
                        cell = tbl.rows[r_idx].cells[c_idx]
                        cell.text = ''
                        para = cell.paragraphs[0]
                        process_inline(para, cell_text)
                        if r_idx == 0:
                            for run in para.runs:
                                run.bold = True
                            set_cell_bg(cell, 'DAEEF3')
                doc.add_paragraph()
            i = next_i
            continue

        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        process_inline(p, stripped)
        i += 1

    doc.save(out_path)
    print(f'  Saved: {out_path}')


# ─── Main ────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_dir = os.path.join(base, 'output')

    slug = sys.argv[1] if len(sys.argv) > 1 else 'active-directory-security-tool'
    slug_dir = os.path.join(output_dir, slug)
    os.makedirs(slug_dir, exist_ok=True)

    print('Building publish .docx ...')
    build_publish_docx(
        os.path.join(slug_dir, f'{slug}_publish.md'),
        os.path.join(slug_dir, f'{slug}_publish.docx')
    )

    print('Building review .docx ...')
    build_review_docx(
        os.path.join(slug_dir, f'{slug}_review.md'),
        os.path.join(slug_dir, f'{slug}_review.docx')
    )

    print('Done.')
